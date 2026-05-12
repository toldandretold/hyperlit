#!/usr/bin/env python3
"""
Regenerate synthetic import-sample fixtures.

Run from repo root:
    pip install python-docx ebooklib
    python3 tests/conversion/import-samples/_generate_synthetic.py

Outputs:
    happy/sample.md
    happy/sample.html
    happy/sample.docx
    happy/sample.epub
    happy/sample.zip
    happy/markdown_folder/{main.md,image.png}
    pdf/peerreview2027.pdf      (tiny valid stub; real OCR side-loaded by the test)
    bad/wrong_ext.txt
    bad/empty.pdf
    bad/oversized.bin           (60 MB — NOT committed; regenerate locally)
"""
from pathlib import Path
import os
import zipfile

from docx import Document        # python-docx
from ebooklib import epub        # ebooklib

ROOT   = Path(__file__).parent
HAPPY  = ROOT / 'happy'
PDFDIR = ROOT / 'pdf'
BAD    = ROOT / 'bad'
for d in (HAPPY, PDFDIR, BAD):
    d.mkdir(exist_ok=True)


# 1. Plain markdown (with one footnote so the conversion code touches its footnote path)
(HAPPY / 'sample.md').write_text(
    "# Sample\n\n"
    "This is a paragraph.\n\n"
    "Another paragraph with a footnote[^1].\n\n"
    "[^1]: A footnote.\n"
)


# 2. Plain HTML
(HAPPY / 'sample.html').write_text(
    "<html><body>"
    "<h1>Sample</h1>"
    "<p>A paragraph.</p>"
    "<p>Another paragraph.</p>"
    "</body></html>"
)


# 3. DOCX
doc = Document()
doc.add_heading('Sample', 0)
doc.add_paragraph('This is a sample document with two paragraphs.')
doc.add_paragraph('And here is the second one.')
doc.save(str(HAPPY / 'sample.docx'))


# 4. EPUB
book = epub.EpubBook()
book.set_identifier('sample-id')
book.set_title('Sample')
book.set_language('en')
ch = epub.EpubHtml(title='Chapter 1', file_name='ch1.xhtml',
                   content='<h1>Sample</h1><p>Body text.</p>')
book.add_item(ch)
book.toc = (ch,)
book.spine = ['nav', ch]
book.add_item(epub.EpubNcx())
book.add_item(epub.EpubNav())
epub.write_epub(str(HAPPY / 'sample.epub'), book)


# 5. Markdown folder + 1×1 PNG
folder = HAPPY / 'markdown_folder'
folder.mkdir(exist_ok=True)
(folder / 'main.md').write_text("# Folder sample\n\n![inline](image.png)\n\nBody text.\n")

# Smallest possible valid PNG: 1×1 transparent
PNG_1x1 = bytes.fromhex(
    '89504e470d0a1a0a'                              # signature
    '0000000d49484452'                              # IHDR length
    '0000000100000001'                              # 1×1
    '0806000000'                                    # bit depth + colour type
    '1f15c489'                                      # IHDR CRC
    '0000000a49444154789c6300010000000500010d0a2db4'   # IDAT
    '0000000049454e44ae426082'                      # IEND
)
(folder / 'image.png').write_bytes(PNG_1x1)


# 6. ZIP of the markdown folder
zip_path = HAPPY / 'sample.zip'
with zipfile.ZipFile(zip_path, 'w', compression=zipfile.ZIP_DEFLATED) as z:
    for p in folder.rglob('*'):
        if p.is_file():
            z.write(p, arcname=p.relative_to(folder.parent))


# 7. Tiny valid PDF stub for the PDF test (real OCR is side-loaded by the test).
#    Built by hand because reportlab isn't installed and the file just needs
#    to pass extension validation + minimal PDF header check.
def build_minimal_pdf() -> bytes:
    objs = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 200 200] /Contents 4 0 R >>",
        b"<< /Length 44 >>\nstream\nBT /F1 12 Tf 50 100 Td (Test stub) Tj ET\nendstream",
    ]
    out = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = []
    for i, body in enumerate(objs, start=1):
        offsets.append(len(out))
        out += f"{i} 0 obj\n".encode() + body + b"\nendobj\n"
    xref_offset = len(out)
    out += b"xref\n0 5\n0000000000 65535 f \n"
    for off in offsets:
        out += f"{off:010d} 00000 n \n".encode()
    out += b"trailer\n<< /Size 5 /Root 1 0 R >>\nstartxref\n"
    out += f"{xref_offset}\n%%EOF\n".encode()
    return bytes(out)

stub_bytes = build_minimal_pdf()
# One stub per fixture in pdf/manifest.json — the test side-loads the
# matching ocr_response.json so the stub bytes are never OCR'd in test.
(PDFDIR / 'whole_document_example.pdf').write_bytes(stub_bytes)
(PDFDIR / 'stem_bibliography_example.pdf').write_bytes(stub_bytes)


# 8. Bad: wrong extension (rejected by extension validator)
(BAD / 'wrong_ext.txt').write_text("plain text rejected by extension validator")

# 9. Bad: empty pdf (rejected by content validator)
(BAD / 'empty.pdf').write_bytes(b'')

# 10. Bad: oversized 60 MB blob with valid extension (so it hits the SIZE check,
#     not the extension check). NOT committed — regenerate locally before running tests.
with open(BAD / 'oversized.pdf', 'wb') as f:
    f.write(os.urandom(60 * 1024 * 1024))


print("Synthetic fixtures written under tests/conversion/import-samples/")
print("  happy/   : sample.{md,html,docx,epub,zip} + markdown_folder/")
print("  pdf/     : whole_document_example.pdf, stem_bibliography_example.pdf (stubs)")
print("  bad/     : wrong_ext.txt, empty.pdf, oversized.pdf (60 MB, gitignored)")
