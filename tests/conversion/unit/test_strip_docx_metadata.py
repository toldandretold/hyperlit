"""Unit tests for the docx front-end — strip_docx_metadata.py.

NOTE ON SCOPE: this module is privacy metadata-scrubbing ONLY. The Word *footnote/citation*
logic does not live here — .docx goes pandoc -> html -> the shared core, so Word footnote
behaviour (incl. the multi-paragraph / blank-line-between-definitions rule) is covered by
test_footnote_extraction.py against process_whole_document_footnotes. These tests just pin
that author/identifying metadata is actually removed.
"""

from docx import Document

from strip_docx_metadata import strip_metadata


def _make_docx(path, author='Secret Author', title='Confidential Title'):
    doc = Document()
    doc.add_paragraph('Body text of the document.')
    props = doc.core_properties
    props.author = author
    props.title = title
    props.last_modified_by = author
    doc.save(str(path))
    return path


def test_strips_author_and_core_properties(tmp_path):
    src = _make_docx(tmp_path / 'in.docx')
    out = tmp_path / 'out.docx'

    assert strip_metadata(str(src), str(out)) is True

    props = Document(str(out)).core_properties
    assert not props.author or props.author in ('', 'None')
    assert not props.last_modified_by
    # body content is preserved
    assert 'Body text of the document.' in '\n'.join(p.text for p in Document(str(out)).paragraphs)


def test_missing_file_returns_false(tmp_path):
    assert strip_metadata(str(tmp_path / 'does_not_exist.docx')) is False


def test_legacy_doc_extension_rejected(tmp_path):
    legacy = tmp_path / 'old.doc'
    legacy.write_bytes(b'not really a doc')
    assert strip_metadata(str(legacy)) is False
